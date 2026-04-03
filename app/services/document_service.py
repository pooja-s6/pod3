"""Document analysis and extraction service."""

import os
import tempfile
from typing import Optional, Dict, List
import uuid
from datetime import datetime

# PDF parsing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# DOCX parsing
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Image text extraction (OCR)
try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None


class DocumentService:
    """Service for handling document upload, parsing, and analysis."""
    
    # Store documents in memory (in production, use database)
    _documents = {}
    
    def __init__(self):
        self.supported_formats = {
            'pdf': '.pdf',
            'docx': '.docx',
            'doc': '.doc',
            'jpg': '.jpg',
            'jpeg': '.jpeg',
            'png': '.png',
            'txt': '.txt'
        }
    
    def upload_and_extract(self, file_path: str, filename: str, session_id: str) -> Dict:
        """
        Upload document and extract text content.
        
        Args:
            file_path: Path to uploaded file
            filename: Original filename
            session_id: Session ID for associating document
        
        Returns:
            Document info with extracted content
        """
        try:
            file_ext = os.path.splitext(filename)[1].lower().lstrip('.')
            
            # Validate file type
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Extract content based on file type
            if file_ext == 'pdf':
                content = self._extract_pdf(file_path)
            elif file_ext == 'docx':
                content = self._extract_docx(file_path)
            elif file_ext in ['txt']:
                content = self._extract_txt(file_path)
            elif file_ext in ['jpg', 'jpeg', 'png']:
                content = self._extract_image_text(file_path)
            else:
                raise ValueError(f"File type {file_ext} extraction not implemented")
            
            # Create document record
            doc_id = str(uuid.uuid4())
            doc_info = {
                'id': doc_id,
                'filename': filename,
                'file_type': file_ext,
                'session_id': session_id,
                'content': content,
                'content_length': len(content),
                'uploaded_at': datetime.now().isoformat(),
                'is_analyzed': True
            }
            
            # Store document
            self._documents[doc_id] = doc_info
            
            return {
                'status': 'success',
                'document_id': doc_id,
                'filename': filename,
                'file_type': file_ext,
                'content_preview': content[:200] + '...' if len(content) > 200 else content,
                'total_characters': len(content),
                'message': 'Document uploaded and analyzed successfully. You can now ask questions about it.'
            }
        
        except Exception as e:
            return {
                'status': 'error',
                'message': str(e)
            }
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        text_content = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
            
            return '\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Failed to extract PDF: {str(e)}")
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if DocxDocument is None:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Failed to extract DOCX: {str(e)}")
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Failed to extract TXT: {str(e)}")
    
    def _extract_image_text(self, file_path: str) -> str:
        """Extract text from image using OCR (pytesseract)."""
        if pytesseract is None or Image is None:
            raise ImportError(
                "pytesseract and Pillow not installed.\n"
                "Run: pip install pytesseract pillow\n"
                "Also install Tesseract-OCR: https://github.com/UB-Mannheim/tesseract/wiki"
            )
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                return "[WARNING] No text detected in image. Image may be unclear or not contain text."
            
            return text
        except Exception as e:
            raise Exception(f"Failed to extract image text: {str(e)}")
    
    def get_document(self, doc_id: str) -> Optional[Dict]:
        """Get document by ID."""
        return self._documents.get(doc_id)
    
    def get_session_documents(self, session_id: str) -> List[Dict]:
        """Get all documents for a session."""
        return [
            {
                'id': doc['id'],
                'filename': doc['filename'],
                'file_type': doc['file_type'],
                'uploaded_at': doc['uploaded_at'],
                'content_length': doc['content_length']
            }
            for doc in self._documents.values()
            if doc['session_id'] == session_id
        ]
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete document by ID."""
        if doc_id in self._documents:
            del self._documents[doc_id]
            return True
        return False
    
    def clear_session_documents(self, session_id: str) -> int:
        """Delete all documents for a session. Returns count deleted."""
        docs_to_delete = [
            doc_id for doc_id, doc in self._documents.items()
            if doc['session_id'] == session_id
        ]
        
        for doc_id in docs_to_delete:
            del self._documents[doc_id]
        
        return len(docs_to_delete)


# Singleton instance
document_service = DocumentService()
